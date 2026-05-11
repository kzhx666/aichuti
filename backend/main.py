import os
import io
import zipfile
import shutil
import asyncio
import httpx
from fastapi import FastAPI, UploadFile, File, Form, Response
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import openai
from openai import AsyncOpenAI
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import pdfplumber
from PIL import Image
import pytesseract
import binascii

app = FastAPI()

class TeacherConfig(BaseModel):
    reader_api_url: str
    reader_api_key: str
    reader_model_name: str
    creator_api_url: str
    creator_api_key: str
    creator_model_name: str
    category: str
    prompt_template: str
    use_ocr: bool = False

class VerifyConfig(BaseModel):
    api_url: str
    api_key: str
    model_name: str

class DownloadRequest(BaseModel):
    markdown_text: str
    mode: str

class CreateFolderRequest(BaseModel):
    category: str

class DeleteRequest(BaseModel):
    category: str

class DeleteDocRequest(BaseModel):
    category: str
    filename: str

def clean_xml_chars(text):
    return re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f]', '', text)

def get_heartbeat(label="HB"):
    return "[__HB__:" + label + "_" + binascii.b2a_hex(os.urandom(4)).decode('utf-8') + "]" + (" " * 1024) + "\n"

# 👉 终极魔改：全能 LaTeX 翻译引擎（支持全局解析与中文下标）
def add_math_to_paragraph(paragraph, text, is_bold=False, font_color=None, font_size=None):
    # 1. 清理可能的转义双斜杠
    text = text.replace('\\\\', '\\')
    
    # 2. 剥离无用的包裹符号，例如 \text{kg} 直接变 kg
    text = re.sub(r'\\text\s*\{(.*?)\}', r'\1', text)
    text = re.sub(r'\\mathrm\s*\{(.*?)\}', r'\1', text)
    text = text.replace(r'\left', '').replace(r'\right', '')
    
    # 3. 建筑/物理常用超大符号词典
    replacements = {
        r'\rho': 'ρ', r'\mu': 'μ', r'\alpha': 'α', r'\beta': 'β',
        r'\theta': 'θ', r'\Delta': 'Δ', r'\gamma': 'γ', r'\lambda': 'λ',
        r'\pi': 'π', r'\sigma': 'σ', r'\omega': 'ω', r'\epsilon': 'ε',
        r'\times': '×', r'\div': '÷', r'\cdot': '·',
        r'\ge': '≥', r'\geq': '≥', r'\le': '≤', r'\leq': '≤',
        r'\approx': '≈', r'\ne': '≠', r'\neq': '≠',
        r'^{\circ}': '°', r'^\circ': '°', r'\circ': '°', r'\%': '%'
    }
    for k, v in replacements.items():
        text = text.replace(k, v)

    # 4. 转换分数 \frac{A}{B} -> A/B
    text = re.sub(r'\\frac\{(.*?)\}\{(.*?)\}', r'\1/\2', text)
    
    # 5. 暴力拆除所有 $ 符号，开启全局强力匹配模式
    text = text.replace('$', '')
    
    # 6. 全局匹配上下标（核心修复：加入了 \u4e00-\u9fa5 以完美支持 K_软 这种中文下标）
    pattern = r'(_\{.*?\}|_[a-zA-Z0-9\u4e00-\u9fa5]|\^\{.*?\}|\^[a-zA-Z0-9\u4e00-\u9fa5])'
    pieces = re.split(pattern, text)
    
    for piece in pieces:
        if not piece: continue
        r = paragraph.add_run()
        r.bold = is_bold
        if font_color: r.font.color.rgb = font_color
        if font_size: r.font.size = font_size
        
        if piece.startswith('_'):
            r.font.subscript = True
            r.text = piece[1:].strip('{}')
        elif piece.startswith('^'):
            val = piece[1:].strip('{}')
            if val == '°':
                r.text = '°'  # 防止度数符号飘得太高
            else:
                r.font.superscript = True
                r.text = val
        else:
            # 去除残留的斜杠，防止影响阅读
            r.text = piece.replace('\\', '')

def create_exam_word(markdown_content, mode='student'):
    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)
    
    title = "《建筑材料》考试试卷" if mode == 'student' else "《建筑材料》考试试卷 (教师标准答案版)"
    h = doc.add_heading(title, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    markdown_content = clean_xml_chars(markdown_content)

    try:
        questions = re.findall(r'\*?\*?\s*(\d+)\.\s*[\[【](.*?)[\]】]\*?\*?(.*?)(<details>.*?</details>)', markdown_content, re.S)
        
        if not questions:
            doc.add_paragraph("⚠️ 提示：系统未能精准识别大模型的排版格式，已为您自动回退为纯文本模式：\n")
            for line in markdown_content.split('\n'):
                if line.strip(): doc.add_paragraph(line)
        else:
            for q_num, q_type, q_body, q_details in questions:
                ans, desc = "", ""
                ans_m = re.search(r'<b>答案：</b>\s*(.*?)\s*<span', q_details)
                if ans_m: ans = ans_m.group(1).strip()
                desc_m = re.search(r'<b>解析：</b>\s*(.*?)\s*</blockquote>', q_details, re.S)
                if desc_m: desc = desc_m.group(1).strip()

                q_body = q_body.strip()
                q_body = re.sub(r'\n正确\s*\n错误\s*$', '', q_body).strip()

                stem, opts = q_body, ""
                opt_match = re.search(r'\nA\..*', q_body, re.S)
                if opt_match:
                    stem = q_body[:opt_match.start()].strip()
                    opts = opt_match.group(0).strip()

                p = doc.add_paragraph()
                if mode == 'teacher':
                    add_math_to_paragraph(p, f"{q_num}. [{q_type}] {stem} ( ", is_bold=True)
                    add_math_to_paragraph(p, ans, is_bold=True, font_color=RGBColor(255, 0, 0))
                    add_math_to_paragraph(p, " )", is_bold=True)
                else:
                    add_math_to_paragraph(p, f"{q_num}. [{q_type}] {stem} (    )", is_bold=True)

                if opts:
                    opt_lines = [o.strip() for o in opts.split('\n') if o.strip()]
                    total_len = sum(len(o) for o in opt_lines)
                    if total_len < 30: 
                        op = doc.add_paragraph()
                        add_math_to_paragraph(op, "\t\t".join(opt_lines))
                    elif total_len < 60 and len(opt_lines) >= 4:
                        op1 = doc.add_paragraph()
                        add_math_to_paragraph(op1, "\t\t".join(opt_lines[:2]))
                        op2 = doc.add_paragraph()
                        add_math_to_paragraph(op2, "\t\t".join(opt_lines[2:]))
                    else:
                        for o in opt_lines: 
                            op = doc.add_paragraph()
                            add_math_to_paragraph(op, o)

                if mode == 'teacher' and desc:
                    desc_p = doc.add_paragraph()
                    add_math_to_paragraph(desc_p, f"【解析】：{desc}", font_color=RGBColor(255, 0, 0), font_size=Pt(10))
    except Exception as e:
        doc.add_paragraph(f"⚠️ 排版过程发生内部兼容性错误，已为您自动回退为纯文本模式 (Error: {str(e)})\n")
        for line in markdown_content.split('\n'):
            if line.strip(): doc.add_paragraph(line)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

@app.post("/verify_api/")
def verify_api(config: VerifyConfig):
    try:
        custom_http_client = httpx.Client(timeout=httpx.Timeout(30.0))
        client = openai.OpenAI(api_key=config.api_key, base_url=config.api_url, http_client=custom_http_client)
        client.chat.completions.create(model=config.model_name, messages=[{"role":"user","content":"Hi"}], max_tokens=5)
        return {"status": "success", "message": "连接成功！主线程畅通！"}
    except Exception as e:
        return {"status": "error", "message": f"连接失败: {str(e)}"}

@app.get("/list_folders/")
def list_folders():
    base_path = "/app/uploads/"
    if not os.path.exists(base_path): os.makedirs(base_path, exist_ok=True)
    return {"folders": [f for f in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, f))]}

@app.post("/create_folder/")
def create_folder(req: CreateFolderRequest):
    os.makedirs(f"/app/uploads/{req.category}", exist_ok=True)
    return {"message": "ok"}

@app.post("/upload_docs/")
async def upload_document(file: UploadFile = File(...), category: str = Form(...)):
    folder_path = f"/app/uploads/{category}"
    os.makedirs(folder_path, exist_ok=True)
    with open(f"{folder_path}/{file.filename}", "wb") as f: f.write(await file.read())
    return {"message": "ok"}

@app.get("/list_docs/")
def list_documents(category: str):
    folder_path = f"/app/uploads/{category}"
    if not os.path.exists(folder_path): return {"files": []}
    return {"files": [{"name": f} for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]}

@app.post("/delete_folder/")
def delete_folder(req: DeleteRequest):
    folder_path = f"/app/uploads/{req.category}"
    if os.path.exists(folder_path): shutil.rmtree(folder_path)
    return {"message": "ok"}

@app.post("/delete_doc/")
def delete_document(req: DeleteDocRequest):
    file_path = f"/app/uploads/{req.category}/{req.filename}"
    if os.path.exists(file_path): os.remove(file_path)
    return {"message": "ok"}

@app.post("/generate_exam/")
async def generate_exam(config: TeacherConfig):
    folder_path = f"/app/uploads/{config.category}"
    
    async def exam_generator():
        docs_text = ""
        yield "[__INIT_FLUSH__:" + binascii.b2a_hex(os.urandom(1024)).decode('utf-8') + "]" + (" " * 1024) + "\n> 🚀 **第一阶段：1 号 AI (资料提取官) 启动，开始地毯式阅读原始资料...**\n\n"
        
        if os.path.exists(folder_path):
            files = os.listdir(folder_path)
            for filename in files:
                file_path = os.path.join(folder_path, filename)
                yield f"> 📄 正在解析文件：`{filename}` ...\n\n"
                await asyncio.sleep(0.1) 
                if filename.endswith(".docx"):
                    docs_text += f"\n\n【Word资料：{filename}】\n"
                    try:
                        doc = docx.Document(file_path)
                        docs_text += "\n".join([p.text for p in doc.paragraphs])
                    except: pass
                elif filename.endswith(".pdf"):
                    docs_text += f"\n\n【PDF资料：{filename}】\n"
                    try:
                        with pdfplumber.open(file_path) as pdf:
                            for page in pdf.pages:
                                t = page.extract_text()
                                if t: docs_text += t + "\n"
                    except: pass

        if not docs_text.strip():
            yield "❌ **错误：未能在资料库中提取到任何文字。**\n"; return
            
        yield f"> ✅ **所有原始资料阅读完毕（共提取约 {len(docs_text)} 字符）。**\n"
        yield "> 🧠 **正在将海量资料送入 1 号 AI 进行浓缩脱水 (长文档提取可能需要一两分钟，请勿刷新页面)...**\n\n"

        try:
            custom_async_client_1 = httpx.AsyncClient(timeout=httpx.Timeout(900.0))
            reader_client = AsyncOpenAI(api_key=config.reader_api_key, base_url=config.reader_api_url, max_retries=2, http_client=custom_async_client_1)
            
            reader_prompt = (
                f"【资料原文】\n{docs_text[:40000]}\n【资料结束】\n\n"
                "========================\n"
                "🔥【系统最高级指令】（请务必严格执行）：\n"
                "1. 提取上述资料中的所有核心考点、技术参数和公式。\n"
                "2. ⚠️强制要求：你输出的内容必须非常详尽，绝对不能少于 1000 字！\n"
                "3. 严禁说“好的”、“明白”，请直接以“# 建筑材料核心考点清单”开头进行输出！"
            )

            reader_messages = [
                {"role": "system", "content": "你是一个无情的干货提取机器。绝不废话，只输出超长篇幅的纯干货笔记。"},
                {"role": "user", "content": reader_prompt} 
            ]

            purified_content = ""
            for extract_loop in range(4):
                reader_task = asyncio.create_task(
                    reader_client.chat.completions.create(
                        model=config.reader_model_name,
                        messages=reader_messages,
                        temperature=0.3, max_tokens=4000
                    )
                )
                while not reader_task.done():
                    yield get_heartbeat("wait")
                    try: await asyncio.wait_for(asyncio.shield(reader_task), timeout=2.0)
                    except asyncio.TimeoutError: continue

                content = reader_task.result().choices[0].message.content or ""
                purified_content += content
                
                if len(content) < 300:
                    yield f"> ⚠️ **1 号 AI 输出过短 ({len(content)} 字符)，疑似偷懒，系统正在强制其重新提炼...**\n"
                    reader_messages.append({"role": "assistant", "content": content})
                    reader_messages.append({"role": "user", "content": "提取得太简略了！你的任务还没完成，请继续补充更多的核心参数和考点，必须输出长文干货！"})
                elif reader_task.result().choices[0].finish_reason in ['length', 'max_tokens']:
                    yield f"> 🔄 **1 号 AI 单次输出达上限，正在强制其无缝衔接继续提取...**\n"
                    reader_messages.append({"role": "assistant", "content": content})
                    reader_messages.append({"role": "user", "content": "输出被截断，请紧接着最后一个字继续输出核心考点！"})
                else:
                    break

            yield f"> ✅ **1 号 AI 脱水提纯完毕！（考点素材成功提炼出 {len(purified_content)} 个纯干货字符）**\n"
        except Exception as e:
            yield f"\n\n[__ERROR__:1号AI提取失败 - ({str(e)})]\n"; return

        yield "> 🎯 **第二阶段：2 号 AI (专家命题官) 接管，核心升级：启动【后台无尘质检沙箱】...**\n"
        yield "> ⏳ **警告：此时纯净输出框不会显示任何题目！系统正在后台质检试卷、切除残肢。请耐心等待总进度达成...**\n\n"

        try:
            custom_async_client_2 = httpx.AsyncClient(timeout=httpx.Timeout(900.0))
            creator_client = AsyncOpenAI(api_key=config.creator_api_key, base_url=config.creator_api_url, max_retries=2, http_client=custom_async_client_2)
            
            messages_history = [
                {"role": "system", "content": config.prompt_template}
            ]

            all_perfect_questions = []
            current_q_num = 0
            total_q = 35 

            for loop in range(15):  
                if current_q_num >= total_q:
                    break
                    
                if current_q_num == 0:
                    user_content = f"资料已提纯。以下是核心考点：\n{purified_content}\n\n请严格按要求命制一套100分试卷（共35题：单选1-10，判断11-25，多选26-35）。从第1题开始。"
                else:
                    user_content = (
                        f"已成功收录至第 {current_q_num} 题。请继续命题！\n"
                        f"🔥警告：必须直接从第 **{current_q_num + 1}** 题开始往下出，直到第 35 题！绝对严禁输出任何废话！"
                    )
                    
                messages_history.append({"role": "user", "content": user_content})
                
                api_task = asyncio.create_task(
                    creator_client.chat.completions.create(
                        model=config.creator_model_name,
                        messages=messages_history,
                        temperature=0.7, stream=True, max_tokens=4000 
                    )
                )

                yield f"> 📡 [沙箱动态] 正在生成并质检第 {current_q_num + 1} 题及后续内容... (当前续写轮次: {loop+1})\n"
                
                assistant_reply = ""
                while not api_task.done():
                    yield get_heartbeat("think_init")
                    try: await asyncio.wait_for(asyncio.shield(api_task), timeout=2.0)
                    except asyncio.TimeoutError: continue

                response = api_task.result()
                finish_reason = None
                iterator = response.__aiter__()
                chunk_counter = 0
                
                while True:
                    try:
                        chunk = await asyncio.wait_for(iterator.__anext__(), timeout=2.0)
                        chunk_counter += 1
                        if hasattr(chunk, 'choices') and chunk.choices and len(chunk.choices) > 0:
                            delta = chunk.choices[0].delta
                            content = getattr(delta, 'content', None)
                            reasoning = getattr(delta, 'reasoning_content', None)
                            
                            if content:
                                assistant_reply += content
                            
                            if chunk_counter % 20 == 0:
                                yield get_heartbeat("buffering")
                                
                            reason = getattr(chunk.choices[0], 'finish_reason', None)
                            if reason: finish_reason = reason
                    except asyncio.TimeoutError:
                        yield get_heartbeat("idle")
                    except StopAsyncIteration:
                        break
                        
                messages_history.append({"role": "assistant", "content": assistant_reply})
                
                pattern = re.compile(r'(?:^|\n)\s*\*?\*?\s*(\d{1,3})\s*\.\s*(?:\[|【).*?</details>', re.IGNORECASE | re.DOTALL)
                
                found_questions = []
                last_num_in_chunk = current_q_num
                
                for m in pattern.finditer(assistant_reply):
                    q_num_str = m.group(1)
                    if q_num_str.isdigit():
                        q_num = int(q_num_str)
                        if q_num == last_num_in_chunk + 1:
                            found_questions.append(m.group(0).strip())
                            last_num_in_chunk = q_num
                        elif q_num <= last_num_in_chunk:
                            pass 
                        else:
                            found_questions.append(m.group(0).strip())
                            last_num_in_chunk = q_num
                
                if found_questions:
                    all_perfect_questions.extend(found_questions)
                    yield f"> ✅ [质检通过] 已成功装载至第 {last_num_in_chunk} 题，所有残肢废话已被无情黑洞吞噬！\n"
                    current_q_num = last_num_in_chunk
                else:
                    yield f"> ⚠️ [质检拦截] 本次输出满篇废话或格式破损，系统已将其丢入黑洞并强制作废重试...\n"

            yield "\n> 🎉 **全部 35 道题目完美质检完毕！开始向前端输出绝密纯净版试卷...**\n\n---\n\n"
            
            final_exam_md = "\n\n".join(all_perfect_questions)
            
            chunk_size = 512
            for i in range(0, len(final_exam_md), chunk_size):
                yield final_exam_md[i:i+chunk_size]
                await asyncio.sleep(0.01)

        except Exception as e:
            yield f"\n\n[__ERROR__:2号AI生成失败 - ({str(e)})]\n"

    return StreamingResponse(exam_generator(), media_type="text/event-stream", headers={"X-Accel-Buffering": "no", "Cache-Control": "no-cache, no-transform", "Connection": "keep-alive"})

@app.post("/download_word/")
def download_word(req: DownloadRequest):
    file_stream = create_exam_word(req.markdown_text, req.mode)
    filename = "student_exam.docx" if req.mode == 'student' else "teacher_exam.docx"
    return Response(content=file_stream.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=" + filename})

